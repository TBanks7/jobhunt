# =============================================================================
# generator.py — Claude tailors LaTeX resume + DOCX cover letter per job
# =============================================================================

import copy
import logging
import os
import re
import shutil
from datetime import datetime
from pathlib import Path

import anthropic
from docx import Document

from config import (
    ANTHROPIC_API_KEY, CLAUDE_MODEL, CANDIDATE_PROFILE,
    RESUME_TEX, COVER_LETTER_DOC, OUTPUT_DIR
)

log = logging.getLogger(__name__)
client = anthropic.Anthropic(api_key=ANTHROPIC_API_KEY)


# ---------------------------------------------------------------------------
# Output directory helpers
# ---------------------------------------------------------------------------

def _make_output_dir(company: str, title: str) -> Path:
    """Create and return output/{company}_{title}/ directory."""
    safe = re.sub(r'[^\w\s-]', '', f"{company}_{title}").strip().replace(' ', '_')
    safe = safe[:60]  # cap length
    timestamp = datetime.now().strftime("%Y%m%d_%H%M")
    folder = Path(OUTPUT_DIR) / f"{safe}_{timestamp}"
    folder.mkdir(parents=True, exist_ok=True)
    return folder


# ---------------------------------------------------------------------------
# RESUME — LaTeX tailoring
# ---------------------------------------------------------------------------

RESUME_SYSTEM_PROMPT = """
You are an expert technical resume editor specializing in software engineering roles.
You will receive:
1. A LaTeX resume source
2. A job description
3. The candidate's background

Your job is to edit the LaTeX resume to maximize relevance for this specific role.

RULES — what you MAY do:
- Reword bullet points to match job description language exactly
- Reorder bullet points within a role (most relevant first)
- Change role titles to better reflect the job (e.g. "Full-Stack Engineer" → "Data Engineer") 
  BUT only if the actual work done genuinely supports the new title
- Add realistic, conservative metrics if they are missing (%, time saved, scale, reliability)
- Strengthen action verbs (built, implemented, optimized, automated, designed, etc.)
- Edit the OBJECTIVE section to match the role perfectly
- Edit the TECHNICAL SKILLS section to mirror keywords in the job description
- Reorder skill categories to surface the most relevant ones first

The progression story the resume should tell:

- Cloud Interactive (2020): Learning fundamentals, executing under guidance
- CSDC (2021–2023): Growing from intern to a capable junior contributor — solid execution, increasing independence, no leadership claims
- Bennett (2022–2024): First taste of ownership — leading small engagements, making technical decisions, coordinating with clients
- Migranium (2024–2025): Confident engineer — mentoring, architecture input, driving quality

If role is more junior than intermediate, use just CSDC experience and Migranium experience and downplay Migranium seniority.

RULES — what you MUST NOT do:
- Do not invent experiences, projects, or companies that don't exist
- Do not change employment dates
- Do not fabricate degrees or certifications
- Do not break LaTeX syntax — output must compile cleanly
- Do not add any explanation or commentary — output ONLY the complete .tex file

METRICS GUIDELINES:
- Performance: "reduced latency by 25%", "improved query performance by 40%"
- Productivity: "cut manual effort by 35%", "automated 80% of reporting workflow"
- Scale: "served 5k+ users", "processed 1M+ records daily"
- Reliability: "maintained 99.9% uptime"
- Delivery: "reduced deployment time by 30%"
Keep metrics believable and conservative.

OUTPUT FORMAT:
Return ONLY the complete, valid LaTeX source. No markdown, no explanation, no backticks.
"""


def tailor_resume(job: dict) -> tuple[str, str]:
    """
    Use Claude to edit the LaTeX resume for a specific job.
    Returns (tailored_tex_content, keyword_match_report).
    """
    with open(RESUME_TEX, "r", encoding="utf-8") as f:
        base_tex = f.read()

    jd = job.get("description", "") or "No description available."
    title = job.get("title", "")
    company = job.get("company", "")

    user_prompt = f"""
Job Title: {title}
Company: {company}
Location: {job.get('location', '')}

--- JOB DESCRIPTION START ---
{jd[:6000]}
--- JOB DESCRIPTION END ---

--- CANDIDATE PROFILE ---
{CANDIDATE_PROFILE}

--- CURRENT LATEX RESUME ---
{base_tex}

Please tailor the LaTeX resume above for this specific role.
After the LaTeX, on a NEW LINE write exactly:
===KEYWORD_REPORT===
Then list the top 10 keywords from the job description and where each appears in the resume.
Format each line as: keyword → section (or "not present")
"""

    log.info(f"Calling Claude to tailor resume for {company} — {title}")
    response = client.messages.create(
        model=CLAUDE_MODEL,
        max_tokens=4096,
        messages=[{"role": "user", "content": user_prompt}],
        system=RESUME_SYSTEM_PROMPT,
    )

    full_output = response.content[0].text

    # Split on keyword report delimiter
    if "===KEYWORD_REPORT===" in full_output:
        tex_part, kw_part = full_output.split("===KEYWORD_REPORT===", 1)
    else:
        tex_part = full_output
        kw_part = "Keyword report not generated."

    return tex_part.strip(), kw_part.strip()


# ---------------------------------------------------------------------------
# COVER LETTER — DOCX tailoring
# ---------------------------------------------------------------------------

COVER_LETTER_SYSTEM_PROMPT = """
You are an expert cover letter writer for software engineering and data roles.
Write a compelling, concise cover letter (3–4 paragraphs) for the candidate below.

TONE: Professional but personable. Confident. Not generic.
LENGTH: 250–350 words maximum.
FORMAT: Plain paragraphs only — no bullet points, no headers, no salutation/sign-off 
        (the template handles those).

STRUCTURE:
Paragraph 1 — Hook: Why this role + company specifically. Show you know them.
Paragraph 2 — What you bring: 2–3 concrete examples tied directly to the job requirements.
              Use specific tech, metrics, and outcomes from the candidate's background.
Paragraph 3 — Forward-looking: What you'd contribute in the first 90 days.
              Optionally a closing sentence expressing enthusiasm for a conversation.

DO NOT:
- Use clichés like "I am writing to express my interest..."
- Repeat the resume bullet points verbatim
- Be generic — reference the company and role specifically
- Output anything other than the cover letter body paragraphs
- Do not include em-dashes in the output
"""


def tailor_cover_letter(job: dict) -> str:
    """
    Use Claude to generate cover letter body text for a specific job.
    Returns the cover letter body as a plain string.
    """
    jd = job.get("description", "") or "No description available."
    title = job.get("title", "")
    company = job.get("company", "")

    user_prompt = f"""
Job Title: {title}
Company: {company}
Location: {job.get('location', '')}

--- JOB DESCRIPTION ---
{jd[:5000]}

--- CANDIDATE PROFILE ---
{CANDIDATE_PROFILE}

Write the cover letter body paragraphs now.
"""

    log.info(f"Calling Claude to write cover letter for {company} — {title}")
    response = client.messages.create(
        model=CLAUDE_MODEL,
        max_tokens=1024,
        messages=[{"role": "user", "content": user_prompt}],
        system=COVER_LETTER_SYSTEM_PROMPT,
    )

    return response.content[0].text.strip()


# ---------------------------------------------------------------------------
# Write files
# ---------------------------------------------------------------------------

def write_tailored_resume(tex_content: str, output_dir: Path) -> Path:
    """Save the tailored .tex file and compile it to PDF."""
    from compiler import compile_latex_to_pdf

    tex_path = output_dir / "resume.tex"
    with open(tex_path, "w", encoding="utf-8") as f:
        f.write(tex_content)
    log.info(f"Tailored resume .tex saved: {tex_path}")

    pdf_path_str = compile_latex_to_pdf(str(tex_path))
    if pdf_path_str:
        log.info(f"Tailored resume PDF compiled: {pdf_path_str}")
    else:
        log.warning("PDF compilation failed — .tex file is still saved, compile manually.")

    return tex_path


def write_keyword_report(report: str, output_dir: Path) -> Path:
    """Save the keyword match report as a text file."""
    rpt_path = output_dir / "keyword_report.txt"
    with open(rpt_path, "w", encoding="utf-8") as f:
        f.write(report)
    log.info(f"Keyword report saved: {rpt_path}")
    return rpt_path


def write_cover_letter_docx(
    cover_letter_body: str,
    job: dict,
    output_dir: Path
) -> Path:
    """
    Fill the cover letter DOCX template with the generated body text.
    Looks for {{COVER_LETTER_BODY}}, {{COMPANY}}, {{ROLE}}, {{DATE}} placeholders.
    If template doesn't have placeholders, appends the body after existing content.
    """
    docx_out = output_dir / "cover_letter.docx"

    if os.path.exists(COVER_LETTER_DOC):
        shutil.copy(COVER_LETTER_DOC, docx_out)
        doc = Document(str(docx_out))

        today = datetime.now().strftime("%B %d, %Y")
        replacements = {
            "{{COVER_LETTER_BODY}}": cover_letter_body,
            "{{COMPANY}}":           job.get("company", ""),
            "{{ROLE}}":              job.get("title", ""),
            "{{DATE}}":              today,
        }

        replaced = False
        for para in doc.paragraphs:
            for placeholder, value in replacements.items():
                if placeholder in para.text:
                    # Preserve formatting by replacing runs
                    for run in para.runs:
                        if placeholder in run.text:
                            run.text = run.text.replace(placeholder, value)
                    replaced = True

        # Also check tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        for placeholder, value in replacements.items():
                            if placeholder in para.text:
                                for run in para.runs:
                                    if placeholder in run.text:
                                        run.text = run.text.replace(placeholder, value)
                                replaced = True

        if not replaced:
            # No placeholders found — append body at end of document
            log.warning("No placeholders found in cover letter template. Appending body.")
            doc.add_paragraph("")
            doc.add_paragraph(cover_letter_body)

        doc.save(str(docx_out))

    else:
        # No template — create a minimal DOCX from scratch
        log.warning(f"Cover letter template not found at {COVER_LETTER_DOC}. Creating from scratch.")
        doc = Document()
        today = datetime.now().strftime("%B %d, %Y")
        doc.add_paragraph(today)
        doc.add_paragraph("")
        doc.add_paragraph(f"Re: {job.get('title', '')} — {job.get('company', '')}")
        doc.add_paragraph("")
        doc.add_paragraph(cover_letter_body)
        doc.add_paragraph("")
        doc.add_paragraph("Sincerely,")
        doc.add_paragraph("Temitope Bankole")
        doc.save(str(docx_out))

    log.info(f"Cover letter DOCX saved: {docx_out}")
    return docx_out


# ---------------------------------------------------------------------------
# Main entry point
# ---------------------------------------------------------------------------

def generate_application_docs(job: dict) -> dict:
    """
    Full pipeline: tailor resume + cover letter for one job.
    Returns dict with paths to all output files.
    """
    company = job.get("company", "Unknown")
    title   = job.get("title", "Unknown")

    output_dir = _make_output_dir(company, title)

    # 1. Tailor resume
    tex_content, keyword_report = tailor_resume(job)
    tex_path = write_tailored_resume(tex_content, output_dir)
    write_keyword_report(keyword_report, output_dir)

    # 2. Tailor cover letter
    cover_letter_body = tailor_cover_letter(job)
    docx_path = write_cover_letter_docx(cover_letter_body, job, output_dir)

    return {
        "output_dir":         str(output_dir),
        "resume_tex":         str(tex_path),
        "cover_letter_docx":  str(docx_path),
        "keyword_report":     str(output_dir / "keyword_report.txt"),
    }